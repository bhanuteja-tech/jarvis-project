"""Generated minimal document fixtures (no real resumes, no real PII).

The PDF builder emits a tiny but structurally valid one-page PDF whose text
pypdf extracts verbatim; the DOCX builder uses python-docx itself.
"""

from __future__ import annotations

import io


def make_pdf(lines: list[str]) -> bytes:
    content_parts = ["BT /F1 12 Tf 72 720 Td 14 TL"]
    for line in lines:
        safe = line.replace("(", "").replace(")", "")
        content_parts.append(f"({safe}) Tj T*")
    content_parts.append("ET")
    stream = " ".join(content_parts).encode("latin-1")

    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n"
        + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    buf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_pos = len(buf)
    buf += b"xref\n0 6\n0000000000 65535 f \n"
    for offset in offsets:
        buf += f"{offset:010d} 00000 n \n".encode()
    buf += (
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_pos).encode()
        + b"\n%%EOF"
    )
    return bytes(buf)


def make_empty_pdf() -> bytes:
    """Structurally valid page with NO text operators (scanned-PDF stand-in)."""
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>",
        b"<< /Length 0 >>\nstream\n\nendstream",
    ]
    buf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(buf)
    buf += b"xref\n0 5\n0000000000 65535 f \n"
    for offset in offsets:
        buf += f"{offset:010d} 00000 n \n".encode()
    buf += (
        b"trailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n"
        + str(xref_pos).encode()
        + b"\n%%EOF"
    )
    return bytes(buf)


def make_docx(paragraphs: list[str], *, heading_first: bool = False) -> bytes:
    import docx

    document = docx.Document()
    for index, paragraph in enumerate(paragraphs):
        if heading_first and index == 0:
            document.add_heading(paragraph, level=1)
        else:
            document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_table(rows: list[list[str]]) -> bytes:
    import docx

    document = docx.Document()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


SAMPLE_RESUME_LINES = [
    "Senior Python Engineer",
    "Skills: python, sql, docker, fastapi",
    "Experience",
    "Engineer at Acme",
    "Jan 2020 - Present",
]
